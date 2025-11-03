"""DOCX converter module for converting HTML/Markdown files to DOCX with CSS styling preservation"""

from pathlib import Path
import tempfile
from typing import Any, Optional

from rich.console import Console

from .html_processor import HTMLProcessor

# Try to import pypandoc, fall back to basic conversion if not available
try:
    import pypandoc

    PYPANDOC_AVAILABLE = True
except ImportError:
    PYPANDOC_AVAILABLE = False

# Try to import python-docx for post-processing
try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

    # Define fallback classes
    class RGBColor:
        def __init__(self, r, g, b):
            self.r = r
            self.g = g
            self.b = b

    class Pt:
        def __init__(self, value):
            self.value = value

    class Inches:
        def __init__(self, value):
            self.value = value

    class WD_ALIGN_PARAGRAPH:  # noqa: N801
        LEFT = 0
        CENTER = 1
        RIGHT = 2
        JUSTIFY = 3

    class WD_STYLE_TYPE:  # noqa: N801
        PARAGRAPH = 1

    class Document:
        def __init__(self, *args, **kwargs):
            pass

        def save(self, path):
            pass

        def add_heading(self, _text, _level):
            return None

        def add_paragraph(self, _text, _style=None):
            return None


# Try to import CSS parsing libraries
try:
    import cssutils

    CSSUTILS_AVAILABLE = True
except ImportError:
    CSSUTILS_AVAILABLE = False

try:
    import tinycss2

    TINYCSS2_AVAILABLE = True
except ImportError:
    TINYCSS2_AVAILABLE = False


class CSSToDocxStyleMapper:
    """Map CSS styles to DOCX style definitions"""

    def __init__(self, console: Optional[Console] = None):
        self.console = console or Console()

    def parse_css_file(self, css_file: Path) -> dict[str, dict[str, Any]]:
        """Parse CSS file and extract style rules"""
        if not css_file.exists():
            return {}

        try:
            with open(css_file, encoding="utf-8") as f:
                css_content = f.read()
            return self._parse_css_content(css_content)
        except Exception as e:
            self.console.print(f"[red]❌ Error parsing CSS file {css_file}: {e}[/red]")
            return {}

    def _parse_css_content(self, css_content: str) -> dict[str, dict[str, Any]]:
        """Parse CSS content and return style mappings"""

        if CSSUTILS_AVAILABLE:
            return self._parse_with_cssutils(css_content)
        elif TINYCSS2_AVAILABLE:
            return self._parse_with_tinycss2(css_content)
        else:
            # Basic regex parsing as fallback
            return self._parse_with_regex(css_content)

    def _parse_with_cssutils(self, css_content: str) -> dict[str, dict[str, Any]]:
        """Parse CSS using cssutils library"""
        styles = {}
        try:
            sheet = cssutils.parseString(css_content)
            for rule in sheet:
                if hasattr(rule, "selectorText") and hasattr(rule, "style"):
                    selector = rule.selectorText
                    style_props = {}
                    for prop in rule.style:
                        style_props[prop.name] = prop.value
                    styles[selector] = style_props
        except Exception as e:
            self.console.print(f"[yellow]⚠️  CSS parsing error: {e}[/yellow]")
        return styles

    def _parse_with_tinycss2(self, css_content: str) -> dict[str, dict[str, Any]]:
        """Parse CSS using tinycss2 library"""
        styles = {}
        try:
            rules = tinycss2.parse_stylesheet(css_content)
            for rule in rules:
                if hasattr(rule, "prelude") and hasattr(rule, "content"):
                    selector = "".join(token.serialize() for token in rule.prelude)
                    style_props = {}
                    for decl in rule.content:
                        if hasattr(decl, "name") and hasattr(decl, "value"):
                            style_props[decl.name] = "".join(token.serialize() for token in decl.value)
                    styles[selector] = style_props
        except Exception as e:
            self.console.print(f"[yellow]⚠️  CSS parsing error: {e}[/yellow]")
        return styles

    def _parse_with_regex(self, css_content: str) -> dict[str, dict[str, Any]]:
        """Basic regex-based CSS parsing as fallback"""
        import re

        styles = {}

        # Simple regex to match CSS rules
        pattern = r"([^{]+)\s*\{\s*([^}]+)\s*\}"
        matches = re.findall(pattern, css_content)

        for selector, properties in matches:
            stripped_selector = selector.strip()
            style_props = {}
            for prop in properties.split(";"):
                if ":" in prop:
                    name, value = prop.split(":", 1)
                    style_props[name.strip()] = value.strip()
            styles[stripped_selector] = style_props

        return styles

    def map_css_to_docx_styles(self, css_styles: dict[str, dict[str, Any]]) -> dict[str, Any]:
        """Convert CSS styles to DOCX-compatible format"""
        docx_styles = {}

        for selector, props in css_styles.items():
            docx_style = {}

            # Font properties
            if "font-family" in props:
                docx_style["font_name"] = props["font-family"].strip("'\"")
            if "font-size" in props:
                try:
                    size = float(props["font-size"].replace("px", "").replace("pt", ""))
                    docx_style["font_size"] = Pt(size)
                except (ValueError, AttributeError):
                    pass
            if "font-weight" in props:
                weight = props["font-weight"].lower()
                docx_style["bold"] = weight in ["bold", "bolder", "700", "800", "900"]
            if "font-style" in props:
                style = props["font-style"].lower()
                docx_style["italic"] = style == "italic"

            # Color properties
            if "color" in props:
                docx_style["color"] = self._parse_color(props["color"])
            if "background-color" in props:
                docx_style["shading"] = self._parse_color(props["background-color"])

            # Text alignment
            if "text-align" in props:
                align = props["text-align"].lower()
                if align == "center":
                    docx_style["alignment"] = WD_ALIGN_PARAGRAPH.CENTER
                elif align == "right":
                    docx_style["alignment"] = WD_ALIGN_PARAGRAPH.RIGHT
                elif align == "justify":
                    docx_style["alignment"] = WD_ALIGN_PARAGRAPH.JUSTIFY
                else:
                    docx_style["alignment"] = WD_ALIGN_PARAGRAPH.LEFT

            # Spacing
            if "margin-top" in props or "margin-bottom" in props:
                docx_style["space_before"] = self._parse_spacing(props.get("margin-top", "0"))
                docx_style["space_after"] = self._parse_spacing(props.get("margin-bottom", "0"))

            if docx_style:
                docx_styles[selector] = docx_style

        return docx_styles

    def _parse_color(self, color_str: str) -> Optional[RGBColor]:
        """Parse CSS color to RGBColor"""
        if not color_str or color_str == "transparent":
            return None

        color_str = color_str.strip()

        # Handle hex colors
        if color_str.startswith("#"):
            hex_color = color_str[1:]
            if len(hex_color) == 3:
                hex_color = "".join([c * 2 for c in hex_color])
            try:
                r = int(hex_color[0:2], 16)
                g = int(hex_color[2:4], 16)
                b = int(hex_color[4:6], 16)
                return RGBColor(r, g, b)
            except (ValueError, IndexError):
                pass

        # Handle rgb() colors
        if color_str.startswith("rgb("):
            try:
                values = color_str[4:-1].split(",")
                r = int(values[0].strip())
                g = int(values[1].strip())
                b = int(values[2].strip())
                return RGBColor(r, g, b)
            except (ValueError, IndexError):
                pass

        # Handle named colors (basic set)
        named_colors = {
            "black": RGBColor(0, 0, 0),
            "white": RGBColor(255, 255, 255),
            "red": RGBColor(255, 0, 0),
            "green": RGBColor(0, 128, 0),
            "blue": RGBColor(0, 0, 255),
            "gray": RGBColor(128, 128, 128),
            "grey": RGBColor(128, 128, 128),
        }

        return named_colors.get(color_str.lower())

    def _parse_spacing(self, spacing_str: str) -> Optional[Pt]:
        """Parse CSS spacing to Pt"""
        if not spacing_str or spacing_str == "0":
            return None

        try:
            value = float(spacing_str.replace("px", "").replace("pt", "").replace("em", ""))
            return Pt(value)
        except (ValueError, AttributeError):
            return None


class DOCXConverter:
    """Convert HTML/MD to DOCX with maximum formatting preservation"""

    def __init__(self, console: Optional[Console] = None):
        """Initialize DOCX converter with optional console for output"""
        self.console = console or Console()
        self.html_processor = HTMLProcessor(console=self.console)
        self.style_mapper = CSSToDocxStyleMapper(console=self.console)
        self.pandoc_available = PYPANDOC_AVAILABLE
        self.python_docx_available = PYTHON_DOCX_AVAILABLE

        if not self.pandoc_available:
            self.console.print("[yellow]⚠️  pypandoc not available. Install with: pip install pypandoc[/yellow]")
        if not self.python_docx_available:
            self.console.print("[yellow]⚠️  python-docx not available. Install with: pip install python-docx[/yellow]")

    def convert_to_docx(
        self,
        input_file: Path,
        output_docx: Path,
        style_template: str = "professional",
        css_file: Optional[Path] = None,
        skip_lint: bool = False,
    ) -> bool:
        """
        Convert a markdown or HTML file to DOCX with optional CSS styling

        Args:
            input_file: Path to the input markdown or HTML file
            output_docx: Path for the output DOCX file
            style_template: Style template name (sensylate, executive, ats, professional)
            css_file: Optional path to a CSS file for styling
            skip_lint: Skip HTML processing and validation

        Returns:
            bool: True if conversion successful, False otherwise
        """
        # Detect file type and route to appropriate method
        file_extension = input_file.suffix.lower()

        if file_extension == ".html":
            return self.convert_html_to_docx(input_file, output_docx, style_template, css_file, skip_lint)
        elif file_extension == ".md":
            return self.convert_markdown_to_docx(input_file, output_docx, style_template, css_file, skip_lint)
        else:
            self.console.print(
                f"[red]❌ Error: Unsupported file type: {file_extension}. Only .md and .html files are supported.[/red]"
            )
            return False

    def convert_html_to_docx(
        self,
        html_file: Path,
        output_docx: Path,
        style_template: str = "professional",
        css_file: Optional[Path] = None,
        skip_lint: bool = False,
    ) -> bool:
        """Convert HTML file to DOCX with styling"""
        try:
            # Step 1: Process HTML with class annotation
            html_content, changes = self._process_html(html_file, skip_lint)

            # Store original HTML for class detection (NEW)
            self._original_html = html_file.read_text(encoding="utf-8")

            # Step 2: Generate reference.docx from CSS template
            reference_docx = self._generate_reference_docx(style_template, css_file)

            # Step 3: Convert HTML to DOCX
            if self.pandoc_available:
                success = self._convert_with_pandoc(html_content, reference_docx, output_docx)
            else:
                success = self._basic_html_to_docx(html_content, output_docx, style_template)

            if success:
                # Step 4: Post-process with python-docx for refinements
                if self.python_docx_available:
                    self._post_process_docx(output_docx, style_template)

                self.console.print(f"[green]✅ DOCX created: {output_docx}[/green]")
                return True
            else:
                return False

        except Exception as e:
            self.console.print(f"[red]❌ HTML to DOCX conversion failed: {e}[/red]")
            return False

    def convert_markdown_to_docx(
        self,
        md_file: Path,
        output_docx: Path,
        style_template: str = "professional",
        css_file: Optional[Path] = None,
        _skip_lint: bool = False,
    ) -> bool:
        """Convert Markdown file to DOCX with styling"""
        try:
            # Read markdown content
            with open(md_file, encoding="utf-8") as f:
                md_content = f.read()

            # Step 2: Generate reference.docx from CSS template
            reference_docx = self._generate_reference_docx(style_template, css_file)

            # Step 3: Convert Markdown to DOCX
            if self.pandoc_available:
                success = self._convert_markdown_with_pandoc(md_content, reference_docx, output_docx)
            else:
                # Convert markdown to HTML first, then to DOCX
                import markdown

                md = markdown.Markdown(extensions=["extra", "codehilite", "toc", "tables"])
                html_content = md.convert(md_content)
                success = self._basic_html_to_docx(html_content, output_docx, style_template)

            if success:
                # Step 4: Post-process with python-docx for refinements
                if self.python_docx_available:
                    self._post_process_docx(output_docx, style_template)

                self.console.print(f"[green]✅ DOCX created: {output_docx}[/green]")
                return True
            else:
                return False

        except Exception as e:
            self.console.print(f"[red]❌ Markdown to DOCX conversion failed: {e}[/red]")
            return False

    def _process_html(self, html_file: Path, skip_lint: bool) -> tuple[str, dict]:
        """Process HTML file with linting, optimization, and class annotation"""
        with open(html_file, encoding="utf-8") as f:
            html_content = f.read()

        if not skip_lint:
            # Use existing HTML processor for consistency
            processed_html, changes = self.html_processor.process_html(html_content)
        else:
            processed_html = html_content
            changes = {}

        # Annotate with page break data attributes (NEW)
        annotated_html = self._annotate_html_with_class_data(processed_html)

        return annotated_html, changes

    def _resolve_style_css(self, style_template: str, css_file: Optional[Path]) -> Optional[Path]:
        """
        Resolve CSS file with format-specific override support

        Priority:
        1. Explicit css_file parameter
        2. Format-specific: {style}_docx.css
        3. Base style: {style}.css

        Args:
            style_template: Style name (ats, professional, executive, sensylate)
            css_file: Optional explicit CSS file path

        Returns:
            Path to CSS file to use
        """
        # Priority 1: Explicit CSS file
        if css_file and css_file.exists():
            self.console.print(f"[dim]📄 Using custom CSS: {css_file.name}[/dim]")
            return css_file

        styles_dir = Path(__file__).parent / "styles"

        # Priority 2: Format-specific CSS (e.g., ats_docx.css)
        format_specific_css = styles_dir / f"{style_template}_docx.css"
        if format_specific_css.exists():
            self.console.print(f"[dim]📄 Using DOCX-specific CSS: {format_specific_css.name}[/dim]")
            return format_specific_css

        # Priority 3: Base style CSS (e.g., ats.css)
        base_css = styles_dir / f"{style_template}.css"
        if base_css.exists():
            self.console.print(f"[dim]📄 Using base CSS: {base_css.name}[/dim]")
            return base_css

        # Fallback: ats.css
        fallback = styles_dir / "ats.css"
        if fallback.exists():
            self.console.print(f"[yellow]⚠️  Style '{style_template}' not found, using ats.css[/yellow]")
            return fallback

        self.console.print(f"[red]❌ No CSS file found for style: {style_template}[/red]")
        return None

    def _generate_reference_docx(self, style_template: str, css_file: Optional[Path]) -> Optional[Path]:
        """Generate reference.docx from CSS template"""
        if not self.python_docx_available:
            return None

        try:
            # Create new document
            doc = Document()

            # Set up basic document properties
            doc.core_properties.title = f"Reference Document - {style_template.title()}"
            doc.core_properties.author = "Applyr DOCX Converter"

            # Resolve CSS with format-specific override
            css_path = self._resolve_style_css(style_template, css_file)
            if not css_path:
                return None

            # Parse CSS and create styles
            if css_path.exists():
                css_styles = self.style_mapper.parse_css_file(css_path)
                docx_styles = self.style_mapper.map_css_to_docx_styles(css_styles)
                self._apply_styles_to_document(doc, docx_styles)

            # Create temporary reference document
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
                temp_path = Path(temp_file.name)
                doc.save(str(temp_path))
                return temp_path

        except Exception as e:
            self.console.print(f"[yellow]⚠️  Could not generate reference.docx: {e}[/yellow]")
            return None

    def _apply_styles_to_document(self, doc: Document, docx_styles: dict[str, Any]) -> None:
        """Apply DOCX styles to document"""
        try:
            # Create custom styles based on CSS
            for selector, style_props in docx_styles.items():
                if selector.startswith("h1") or "heading-1" in selector:
                    self._create_heading_style(doc, "Heading 1", style_props)
                elif selector.startswith("h2") or "heading-2" in selector:
                    self._create_heading_style(doc, "Heading 2", style_props)
                elif selector.startswith("h3") or "heading-3" in selector:
                    self._create_heading_style(doc, "Heading 3", style_props)
                elif "body" in selector or "p" in selector:
                    self._create_paragraph_style(doc, "Normal", style_props)
        except Exception as e:
            self.console.print(f"[yellow]⚠️  Error applying styles: {e}[/yellow]")

    def _create_heading_style(self, doc: Document, style_name: str, style_props: dict[str, Any]) -> None:
        """Create heading style"""
        try:
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            if "font_name" in style_props:
                style.font.name = style_props["font_name"]
            if "font_size" in style_props:
                style.font.size = style_props["font_size"]
            if "bold" in style_props:
                style.font.bold = style_props["bold"]
            if "color" in style_props and style_props["color"]:
                style.font.color.rgb = style_props["color"]
        except Exception:
            pass  # Style might already exist

    def _create_paragraph_style(self, doc: Document, style_name: str, style_props: dict[str, Any]) -> None:
        """Create paragraph style"""
        try:
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            if "font_name" in style_props:
                style.font.name = style_props["font_name"]
            if "font_size" in style_props:
                style.font.size = style_props["font_size"]
            if "bold" in style_props:
                style.font.bold = style_props["bold"]
            if "color" in style_props and style_props["color"]:
                style.font.color.rgb = style_props["color"]
        except Exception:
            pass  # Style might already exist

    def _convert_with_pandoc(self, html_content: str, reference_docx: Optional[Path], output_docx: Path) -> bool:
        """Convert HTML to DOCX using pypandoc"""
        try:
            # Ensure output directory exists
            output_docx.parent.mkdir(parents=True, exist_ok=True)

            # Prepare pandoc options
            options = ["--standalone"]

            if reference_docx and reference_docx.exists():
                options.extend(["--reference-doc", str(reference_docx)])

            # Convert HTML to DOCX
            pypandoc.convert_text(html_content, "docx", format="html", outputfile=str(output_docx), extra_args=options)

            return True

        except Exception as e:
            self.console.print(f"[red]❌ Pandoc conversion failed: {e}[/red]")
            # Try fallback conversion
            return self._basic_html_to_docx(html_content, output_docx, "professional")

    def _convert_markdown_with_pandoc(self, md_content: str, reference_docx: Optional[Path], output_docx: Path) -> bool:
        """Convert Markdown to DOCX using pypandoc"""
        try:
            # Ensure output directory exists
            output_docx.parent.mkdir(parents=True, exist_ok=True)

            # Prepare pandoc options
            options = ["--standalone"]

            if reference_docx and reference_docx.exists():
                options.extend(["--reference-doc", str(reference_docx)])

            # Convert Markdown to DOCX
            pypandoc.convert_text(
                md_content, "docx", format="markdown", outputfile=str(output_docx), extra_args=options
            )

            return True

        except Exception as e:
            self.console.print(f"[red]❌ Pandoc conversion failed: {e}[/red]")
            # Try fallback conversion
            import markdown

            md = markdown.Markdown(extensions=["extra", "codehilite", "toc", "tables"])
            html_content = md.convert(md_content)
            return self._basic_html_to_docx(html_content, output_docx, "professional")

    def _basic_html_to_docx(self, html_content: str, output_docx: Path, style_template: str) -> bool:
        """Fallback: Basic HTML to DOCX using python-docx + BeautifulSoup"""
        if not self.python_docx_available:
            self.console.print("[red]❌ python-docx not available for basic conversion[/red]")
            return False

        try:
            from bs4 import BeautifulSoup

            # Parse HTML
            soup = BeautifulSoup(html_content, "html.parser")

            # Create new document
            doc = Document()

            # Extract and add content
            self._extract_content_to_docx(soup, doc, style_template)

            # Save document
            output_docx.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(output_docx))

            return True

        except Exception as e:
            self.console.print(f"[red]❌ Basic HTML to DOCX conversion failed: {e}[/red]")
            return False

    def _extract_content_to_docx(self, soup, doc: Document, _style_template: str) -> None:
        """Extract content from BeautifulSoup and add to DOCX"""
        # Find main content area
        body = soup.find("body") or soup

        for element in body.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "ul", "ol", "table"]):
            if element.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
                # Add heading
                level = int(element.name[1])
                doc.add_heading(element.get_text(strip=True), level=level)

            elif element.name == "p":
                # Add paragraph
                text = element.get_text(strip=True)
                if text:
                    doc.add_paragraph(text)

            elif element.name in ["ul", "ol"]:
                # Add list
                for li in element.find_all("li"):
                    text = li.get_text(strip=True)
                    if text:
                        doc.add_paragraph(text, style="List Bullet" if element.name == "ul" else "List Number")

    def _post_process_docx(self, docx_path: Path, style_template: str) -> None:
        """Post-process DOCX for final refinements"""
        if not self.python_docx_available:
            return

        try:
            # Load document
            doc = Document(str(docx_path))

            # Apply intelligent page breaks with class detection (NEW)
            self._apply_smart_page_breaks(doc, style_template, getattr(self, "_original_html", None))

            # Apply template-specific refinements
            if style_template == "ats":
                self._apply_ats_optimizations(doc)
            elif style_template == "executive":
                self._apply_executive_optimizations(doc)
            elif style_template == "sensylate":
                self._apply_sensylate_optimizations(doc)
            else:
                self._apply_professional_optimizations(doc)

            # Save document
            doc.save(str(docx_path))

            self.console.print("[dim]✨ Applied intelligent page break controls[/dim]")

        except Exception as e:
            self.console.print(f"[yellow]⚠️  Post-processing failed: {e}[/yellow]")

    def _apply_ats_optimizations(self, doc: Document) -> None:
        """Apply ATS-specific optimizations"""
        # Ensure simple formatting for ATS compatibility
        for paragraph in doc.paragraphs:
            # Remove complex formatting that might confuse ATS
            if paragraph.runs:
                for run in paragraph.runs:
                    # Keep only basic formatting
                    if run.bold is None:
                        run.bold = False
                    if run.italic is None:
                        run.italic = False

    def _apply_executive_optimizations(self, doc: Document) -> None:
        """Apply executive-style optimizations"""
        # Add executive formatting touches

    def _apply_sensylate_optimizations(self, doc: Document) -> None:
        """Apply Sensylate-style optimizations"""
        # Add Sensylate branding touches

    def _apply_professional_optimizations(self, doc: Document) -> None:
        """Apply professional-style optimizations"""
        # Standard professional formatting

    # Page Break Intelligence Methods

    def _is_heading(self, para) -> bool:
        """Check if paragraph is a heading"""
        return para.style.name.startswith("Heading") or para.style.name.startswith("Title")

    def _is_major_heading(self, para) -> bool:
        """Check if paragraph is a major heading (h1/h2 only, not h3/h4)"""
        return para.style and para.style.name in ["Heading 1", "Heading 2", "Title"]

    def _is_list_item(self, para) -> bool:
        """Check if paragraph is a list item"""
        return para.style.name in ["List Bullet", "List Number", "List Paragraph", "List Continue"]

    def _is_body_text(self, para) -> bool:
        """Check if paragraph is body text"""
        return para.style.name in ["Normal", "Body Text", "Body Text 2"]

    def _is_likely_job_title(self, text: str) -> bool:
        """Heuristic to detect job titles"""
        # Common patterns: all caps, ends with role words, etc.
        role_indicators = [
            "Engineer",
            "Developer",
            "Manager",
            "Director",
            "Specialist",
            "Analyst",
            "Lead",
            "Senior",
            "Principal",
            "Architect",
            "Consultant",
            "Coordinator",
            "Supervisor",
        ]
        return any(indicator in text for indicator in role_indicators)

    def _is_likely_company_name(self, para) -> bool:
        """Heuristic to detect company names"""
        # Check for bold formatting, specific patterns
        return para.runs and para.runs[0].bold and not self._is_heading(para)

    def _is_likely_achievement_section(self, text: str) -> bool:
        """Heuristic to detect achievement sections"""
        achievement_keywords = ["achievement", "highlight", "accomplishment", "key result", "deliverable", "success"]
        return any(keyword in text.lower() for keyword in achievement_keywords)

    # Enhanced Class-Based Detection Methods

    def _extract_new_page_markers(self, html_content: str) -> dict:
        """
        Extract ONLY new-page class markers

        Returns dict mapping next sibling text to 'new-page'
        """
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(html_content, "html.parser")
        class_map = {}

        # Find all new-page divs
        new_page_divs = soup.find_all(class_="new-page")

        for div in new_page_divs:
            # Find next sibling element
            next_sibling = div.find_next_sibling()
            if next_sibling:
                text = next_sibling.get_text(strip=True)[:100]
                if text:
                    class_map[text] = "new-page"
                    self.console.print(f"[dim]   new-page → '{text[:50]}...'[/dim]")

        return class_map

    def _extract_article_markers(self, html_content: str) -> dict:
        """
        Extract article block markers to prevent page breaks within articles

        Returns dict mapping article content snippets to 'article-block' with paragraph ranges
        """
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(html_content, "html.parser")
        article_map = {}

        # Find all article tags
        articles = soup.find_all("article")

        self.console.print(f"[dim]🔍 Found {len(articles)} article blocks in HTML[/dim]")

        for idx, article in enumerate(articles):
            # Get all text content within the article
            article_text = article.get_text(strip=True)

            # Get first significant text (first heading or paragraph)
            first_element = None
            for tag in ["h1", "h2", "h3", "h4", "h5", "h6", "p"]:
                first_element = article.find(tag)
                if first_element:
                    break

            if first_element:
                first_text = first_element.get_text(strip=True)[:100]

                # Get last significant text
                all_paragraphs = article.find_all(["p", "li", "h1", "h2", "h3", "h4", "h5", "h6"])
                last_text = None
                if all_paragraphs:
                    last_text = all_paragraphs[-1].get_text(strip=True)[:100]

                # Store article boundaries
                article_map[first_text] = {
                    "type": "article-start",
                    "last_text": last_text,
                    "length": len(article_text),
                    "index": idx,
                }

                self.console.print(f"[dim]   article[{idx}] start → '{first_text[:40]}...'[/dim]")
                if last_text and last_text != first_text:
                    self.console.print(f"[dim]   article[{idx}] end   → '{last_text[:40]}...'[/dim]")

        return article_map

    def _annotate_html_with_class_data(self, html_content: str) -> str:
        """
        Add data-page-break attributes ONLY for new-page class
        """
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(html_content, "html.parser")

        # Only handle new-page class
        new_page_divs = soup.find_all(class_="new-page")
        for div in new_page_divs:
            div["data-page-break"] = "before"

        return str(soup)

    def _apply_smart_page_breaks(self, doc, _style_template: str, html_content: str = None) -> None:
        """
        Apply explicit page breaks and article block protection

        1. Explicit page breaks via new-page class
        2. Keep-together for article blocks (prevent splitting across pages)
        """
        if not html_content:
            self.console.print("[dim]✨ No page break rules applied (no HTML content)[/dim]")
            return

        # Extract new-page markers
        class_map = self._extract_new_page_markers(html_content)

        # Extract article boundaries
        article_map = self._extract_article_markers(html_content)

        if not class_map and not article_map:
            self.console.print("[dim]✨ No page break markers or article blocks found[/dim]")
            return

        paragraphs = list(doc.paragraphs)
        page_breaks_applied = 0
        articles_protected = 0

        # Track if we're inside an article block
        current_article = None
        article_paragraphs = []

        for _i, para in enumerate(paragraphs):
            if not para.text.strip():
                continue

            para_snippet = para.text.strip()[:100]

            # Check for explicit page breaks (new-page class)
            if self._should_have_page_break(para.text, class_map):
                para.paragraph_format.page_break_before = True
                page_breaks_applied += 1
                self.console.print(f"[dim]   ✓ Page break before: '{para.text.strip()[:50]}...'[/dim]")

            # Check if this is the start of an article block
            article_info = self._find_article_start(para_snippet, article_map)
            if article_info:
                # If we're already tracking an article, finalize it first
                if current_article and article_paragraphs:
                    self._apply_keep_together_to_article(article_paragraphs)
                    articles_protected += 1
                    self.console.print(
                        f"[dim]   ✓ Protected article[{current_article['index']}] ({len(article_paragraphs)} paragraphs)[/dim]"
                    )

                # Now start tracking the new article
                current_article = article_info
                article_paragraphs = [para]
                self.console.print(f"[dim]   📄 Article[{article_info['index']}] start detected[/dim]")
            elif current_article:
                # We're inside an article, add to collection
                article_paragraphs.append(para)

                # Check if this is the end of the article
                if self._is_article_end(para_snippet, current_article):
                    # Apply keep_together to all paragraphs in this article
                    self._apply_keep_together_to_article(article_paragraphs)
                    articles_protected += 1
                    self.console.print(
                        f"[dim]   ✓ Protected article[{current_article['index']}] ({len(article_paragraphs)} paragraphs)[/dim]"
                    )

                    # Reset article tracking
                    current_article = None
                    article_paragraphs = []

        # Handle case where article didn't have clear end marker
        if current_article and article_paragraphs:
            self._apply_keep_together_to_article(article_paragraphs)
            articles_protected += 1
            self.console.print(
                f"[dim]   ✓ Protected article[{current_article['index']}] ({len(article_paragraphs)} paragraphs, no end marker)[/dim]"
            )

        self.console.print(
            f"[dim]✨ Applied {page_breaks_applied} page breaks and protected {articles_protected} article blocks[/dim]"
        )

    def _should_have_page_break(self, para_text: str, class_map: dict) -> bool:
        """Check if paragraph text matches any new-page marker"""
        para_snippet = para_text.strip()[:100]

        # Direct match
        if para_snippet in class_map:
            return True

        # Fuzzy match (partial overlap)
        for text_snippet in class_map:
            if (text_snippet in para_snippet or para_snippet in text_snippet) and len(text_snippet) > 10:
                return True

        return False

    def _find_article_start(self, para_snippet: str, article_map: dict) -> dict:
        """
        Check if paragraph matches the start of an article block

        Returns article info dict if match found, None otherwise
        """
        # Direct match
        if para_snippet in article_map:
            return article_map[para_snippet]

        # Fuzzy match (partial overlap)
        for text_snippet, article_info in article_map.items():
            if (
                article_info["type"] == "article-start"
                and (text_snippet in para_snippet or para_snippet in text_snippet)
                and len(text_snippet) > 15
            ):
                return article_info

        return None

    def _is_article_end(self, para_snippet: str, article_info: dict) -> bool:
        """
        Check if paragraph matches the end of the current article block

        Returns True if this is the last paragraph of the article
        """
        if not article_info or "last_text" not in article_info:
            return False

        last_text = article_info["last_text"]
        if not last_text:
            return False

        # Check for match with article's last text
        return (last_text in para_snippet or para_snippet in last_text) and len(last_text) > 15

    def _apply_keep_together_to_article(self, article_paragraphs: list) -> None:
        """
        Apply keep_together formatting to all paragraphs in an article
        to prevent the article from being split across pages
        """
        if not article_paragraphs:
            return

        try:
            for para in article_paragraphs:
                # Set keep_together to prevent page break within this paragraph
                para.paragraph_format.keep_together = True

                # Also set keep_with_next for all but the last paragraph
                # This keeps the article as a cohesive block
                if para != article_paragraphs[-1]:
                    para.paragraph_format.keep_with_next = True
        except Exception as e:
            self.console.print(f"[yellow]⚠️  Could not apply keep_together: {e}[/yellow]")
